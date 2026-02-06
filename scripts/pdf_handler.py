#!/usr/bin/env python3
"""
PDF 转 Word 转换器
使用 pdfplumber 和 python-docx 实现 PDF 到 DOCX 的转换
"""

from pathlib import Path
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def pdf_to_word(input_path: str, output_path: str) -> dict:
    """
    将 PDF 文件转换为 Word 文档
    
    Args:
        input_path: PDF 文件路径
        output_path: 输出 Word 文件路径
    
    Returns:
        dict: 转换结果信息
    """
    result = {
        "success": False,
        "pages": 0,
        "message": ""
    }
    
    try:
        # 创建 Word 文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(Path(input_path).stem, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 打开 PDF
        with pdfplumber.open(input_path) as pdf:
            result["pages"] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                # 添加页面标题
                doc.add_heading(f"第 {page_num} 页", level=1)
                
                # 提取文本
                text = page.extract_text()
                
                if text:
                    # 清理文本
                    text = clean_text(text)
                    
                    # 按段落分割并添加
                    paragraphs = text.split('\n')
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            # 检查是否是标题（短行+无句末标点）
                            if len(para) < 50 and not para.endswith(('。', '！', '？', ')', ']')):
                                heading = doc.add_heading(para, level=2)
                            else:
                                p = doc.add_paragraph(para)
                
                # 提取表格（如果有）
                tables = page.extract_tables()
                if tables:
                    doc.add_heading("表格", level=3)
                    for table in tables:
                        table_doc = doc.add_table(rows=1, cols=len(table[0]) if table else 0)
                        table_doc.style = 'Table Grid'
                        
                        # 添加表头
                        header_cells = table_doc.rows[0].cells
                        for i, cell in enumerate(table[0]):
                            header_cells[i].text = str(cell) if cell else ""
                        
                        # 添加数据行
                        for row in table[1:]:
                            row_cells = table_doc.rows[-1].cells
                            for i, cell in enumerate(row):
                                if i < len(row_cells):
                                    row_cells[i].text = str(cell) if cell else ""
                
                # 添加页面分隔
                if page_num < len(pdf.pages):
                    doc.add_page_break()
        
        # 保存文档
        doc.save(output_path)
        
        result["success"] = True
        result["message"] = f"转换成功！共 {result['pages']} 页"
        
    except Exception as e:
        result["message"] = f"转换失败: {str(e)}"
    
    return result


def clean_text(text: str) -> str:
    """清理提取的文本"""
    # 移除多余的空白字符
    text = re.sub(r'\s+', ' ', text)
    # 修复常见的 PDF 提取问题
    text = text.replace('\x00', '')
    return text.strip()


def convert_batch(input_dir: str, output_dir: str) -> list:
    """
    批量转换 PDF 文件
    
    Args:
        input_dir: 输入文件夹路径
        output_dir: 输出文件夹路径
    
    Returns:
        list: 每个文件的转换结果列表
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    
    results = []
    pdf_files = list(input_path.glob("*.pdf"))
    
    for pdf_file in pdf_files:
        output_file = output_path / f"{pdf_file.stem}.docx"
        result = pdf_to_word(str(pdf_file), str(output_file))
        result["file"] = pdf_file.name
        results.append(result)
    
    return results


if __name__ == "__main__":
    # 测试转换
    import sys
    
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.pdf', '.docx')
        
        result = pdf_to_word(input_file, output_file)
        print(result)
    else:
        print("用法: python pdf_handler.py <input.pdf> [output.docx]")
